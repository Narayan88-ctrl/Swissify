from __future__ import annotations

# ---------- stdlib ----------
from pathlib import Path
from io import BytesIO
import json
from typing import Optional
import re

# ---------- third-party ----------
from fastapi import FastAPI, Request, Form, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import httpx

# ---------- paths & constants ----------
ROOT = Path(__file__).resolve().parent
TEMPLATES_DIR = ROOT / "templates"
STATIC_DIR = ROOT / "static"
I18N_DIR = ROOT / "i18n"

ENGINE_BASE = "http://127.0.0.1:8600"   # Swissify Engine

ALLOWED_LANGS = {"en", "de"}           # EN/DE only for demo
DEFAULT_LANG = "en"

# ---------- app ----------
app = FastAPI(title="Swissify Web", version="3.0.0")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))


# ---------- i18n helper ----------
def load_lang(lang: str) -> dict:
    if lang not in ALLOWED_LANGS:
        lang = DEFAULT_LANG
    try:
        data = json.loads((I18N_DIR / f"{lang}.json").read_text(encoding="utf-8"))
        if isinstance(data, dict):
            return data
    except Exception:
        pass
    # minimal fallback
    if lang == "de":
        return {
            "brand": "SWISSIFY ENGINE",
            "swiss_standard": "Swiss Standard",
            "language": "Sprache",
            "teacher_login": "Lehrer-Login",
            "input_title": "Swissify Text",
            "input_hint": "Füge beliebigen deutschen Text ein; wir wandeln in den Schweizer Standard um (Zahlen, Anführungszeichen, Lexikon, Orthografie).",
            "convert_now": "Jetzt konvertieren",
            "clear": "Leeren",
            "output_title": "Ausgabe",
            "output_hint": "Konvertierter Schweizer Standard.",
            "download_txt": ".txt herunterladen",
            "docx_title": "DOCX-Export",
            "docx_hint": "*.txt / *.docx hochladen → Swissified DOCX herunterladen.",
            "pdf_title": "PDF-Export",
            "pdf_hint": "*.txt oder echtes *.pdf hochladen → Swissified PDF herunterladen.",
            "footer_left": "SWISSIFY ENGINE",
            "footer_right": "Engine: FastAPI · Web: Jinja",
            "upload": "Datei wählen",
            "convert_docx": "DOCX umwandeln",
            "convert_pdf": "PDF umwandeln",
            "error_bad_file": "Bitte eine echte Datei mit der richtigen Endung hochladen.",
            "error_engine": "Engine nicht erreichbar. Bitte prüfen Sie Port 8600 (/health).",
        }
    else:
        return {
            "brand": "SWISSIFY ENGINE",
            "swiss_standard": "Swiss Standard",
            "language": "Language",
            "teacher_login": "Teacher Login",
            "input_title": "Swissify Text",
            "input_hint": "Paste any German text; we convert to Swiss standard (numbers, quotes, lexicon, orthography).",
            "convert_now": "Convert Now",
            "clear": "Clear",
            "output_title": "Output",
            "output_hint": "Converted Swiss standard text.",
            "download_txt": "Download .txt",
            "docx_title": "DOCX Export",
            "docx_hint": "Upload *.txt / *.docx → Download Swissified DOCX.",
            "pdf_title": "PDF Export",
            "pdf_hint": "Upload *.txt or real *.pdf → Download Swissified PDF.",
            "footer_left": "SWISSIFY ENGINE",
            "footer_right": "Engine: FastAPI · Web: Jinja",
            "upload": "Choose file",
            "convert_docx": "Convert DOCX",
            "convert_pdf": "Convert PDF",
            "error_bad_file": "Please upload a real file with correct extension.",
            "error_engine": "Engine not reachable. Check port 8600 (/health).",
        }


def render(request: Request, name: str, lang: str, ctx: dict | None = None):
    t = load_lang(lang)
    base = {"request": request, "lang": lang, "t": t}
    if ctx:
        base.update(ctx)
    return templates.TemplateResponse(name, base)


# ---------- engine call ----------
async def swissify_text(raw: str) -> str:
    try:
        async with httpx.AsyncClient(timeout=15.0) as cl:
            r = await cl.post(f"{ENGINE_BASE}/convert", json={"text": raw})
            r.raise_for_status()
            data = r.json()
            return data.get("output", "")
    except Exception:
        return ""


# ---------- routes ----------
@app.get("/", response_class=HTMLResponse)
async def home(request: Request, lang: str = DEFAULT_LANG):
    return render(request, "index.html", lang, {"text": "", "output": ""})


@app.post("/convert_text", response_class=HTMLResponse)
async def convert_text(request: Request, lang: str = DEFAULT_LANG, text: str = Form("")):
    out = await swissify_text(text or "")
    return render(request, "index.html", lang, {"text": text, "output": out})


# --- Teacher login (demo/dummy) ---
@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request, lang: str = DEFAULT_LANG):
    return render(request, "login.html", lang, {})

@app.post("/login")
async def login_submit(request: Request, lang: str = DEFAULT_LANG, username: str = Form(""), password: str = Form("")):
    # Demo only: accept anything and go home
    return RedirectResponse(url=f"/?lang={lang}", status_code=303)


# ---------- file helpers ----------
def _ext_ok(filename: str, allowed: tuple[str, ...]) -> bool:
    return bool(filename) and filename.lower().endswith(allowed)


# --- DOCX export ---
@app.post("/convert_docx", response_class=StreamingResponse)
async def convert_docx(lang: str = DEFAULT_LANG, file: UploadFile = File(...)):
    if not _ext_ok(file.filename or "", (".txt", ".docx")):
        # clean message
        raise ValueError(load_lang(lang)["error_bad_file"])

    text = ""
    name = Path(file.filename or "input").stem

    if file.filename.lower().endswith(".txt"):
        text = (await file.read()).decode("utf-8", errors="replace")
    else:
        from docx import Document  # python-docx
        doc = Document(BytesIO(await file.read()))
        text = "\n".join(p.text for p in doc.paragraphs)

    out = await swissify_text(text)
    from docx import Document
    from docx.shared import Pt

    d = Document()
    p = d.add_paragraph(out.strip())
    p.style.font.name = "Arial"
    p.style.font.size = Pt(11)
    bio = BytesIO()
    d.save(bio)
    bio.seek(0)

    headers = {"Content-Disposition": f'attachment; filename="{name}_swissified.docx"'}
    return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


# --- PDF export ---
@app.post("/convert_pdf", response_class=StreamingResponse)
async def convert_pdf(lang: str = DEFAULT_LANG, file: UploadFile = File(...)):
    if not _ext_ok(file.filename or "", (".txt", ".pdf")):
        raise ValueError(load_lang(lang)["error_bad_file"])

    name = Path(file.filename or "input").stem
    text = ""

    if file.filename.lower().endswith(".txt"):
        text = (await file.read()).decode("utf-8", errors="replace")
    else:
        # Try text extraction; if fails we still Swissify per page text payload
        try:
            from pdfminer.high_level import extract_text
            raw_bytes = await file.read()
            text = extract_text(BytesIO(raw_bytes))
        except Exception:
            text = ""  # empty → swissify returns ""

    out = await swissify_text(text)

    # Build a clean PDF using reportlab
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    x, y = 20 * mm, height - 25 * mm
    c.setFont("Helvetica", 11)

    for line in (out or "").splitlines() or ["(empty)"]:
        if y < 20 * mm:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 25 * mm
        c.drawString(x, y, line)
        y -= 6 * mm

    c.showPage()
    c.save()
    bio.seek(0)

    headers = {"Content-Disposition": f'attachment; filename="{name}_swissified.pdf"'}
    return StreamingResponse(bio, media_type="application/pdf", headers=headers)
